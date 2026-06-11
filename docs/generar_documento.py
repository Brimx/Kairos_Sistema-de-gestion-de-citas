from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Estilos base ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Helper para títulos sin estilo heading ──
def add_title(text, size=14):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

# ── Configurar página ──
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ── Portada ──
for _ in range(6):
    doc.add_paragraph()

titulo = doc.add_paragraph()
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = titulo.add_run('Kairos — Sistema de Gestión de Citas Médicas')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

subt = doc.add_paragraph()
subt.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subt.add_run('Documentación del Proyecto')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(
    'Curso: Programación Orientada a Objetos\n'
    'Docente: Patricia Isabel Álvarez Ortega\n'
    'Universidad Popular del Cesar\n'
    'Facultad de Ingenierías y Tecnológicas\n\n'
    'Estudiantes:\n'
    '• Santiago Andrés Montenegro Muñoz\n'
    '• Kevin Jhosep Mercado Morón\n'
    '• Jose Daniel Pinzon Racero\n\n'
    'Junio 2026'
)
run.font.size = Pt(12)

doc.add_page_break()

# ── Índice ──
add_title('Índice')
indice_items = [
    '1. Planteamiento del problema',
    '2. Objetivo general y objetivos específicos',
    '   2.1 Objetivo general',
    '   2.2 Objetivos específicos',
    '3. Justificación',
    '4. Diagrama de clases UML',
    '5. Manual de usuario',
    '   5.1 Requisitos del sistema',
    '   5.2 Instalación y ejecución',
    '   5.3 Inicio de sesión',
    '   5.4 Registro de paciente',
    '   5.5 Dashboard del paciente',
    '   5.6 Dashboard del médico',
    '   5.7 Dashboard del administrador',
    '   5.8 Funcionalidades comunes',
]
for item in indice_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ════════════════════════════════════════════
# 1. PLANTEAMIENTO DEL PROBLEMA
# ════════════════════════════════════════════
add_title('1. Planteamiento del problema')

doc.add_paragraph(
    'En la mayoría de las instituciones prestadoras de salud (IPS) de la región, '
    'el proceso de agendamiento de citas médicas se maneja de forma manual o '
    'semimanual: el paciente debe llamar por teléfono o asistir presencialmente '
    'a la recepción para solicitar una cita, el personal administrativo consulta '
    'una agenda física o un archivo de Excel para verificar disponibilidad, y '
    'finalmente asigna un horario de forma verbal o mediante anotación en una '
    'libreta.'
)

doc.add_paragraph(
    'Este flujo presenta múltiples inconvenientes:'
)

problemas = [
    'Pérdida de información: las anotaciones manuales son propensas a errores, '
    'tachaduras y pérdida de hojas, lo que puede resultar en la pérdida total '
    'del registro de citas.',
    'Doble agendamiento: al no existir un sistema centralizado en tiempo real, '
    'es frecuente que dos pacientes sean agendados en el mismo horario con el '
    'mismo médico.',
    'Falta de trazabilidad: no queda registro histórico de las citas anteriores '
    'del paciente, sus diagnósticos, recetas ni tratamientos previos.',
    'Inasistencias no controladas: no existe un mecanismo automático para '
    'detectar y registrar las citas a las que el paciente no asistió, lo que '
    'dificulta la toma de decisiones sobre sobrecupo o reprogramación.',
    'Dificultad para gestionar horarios: los médicos no tienen visibilidad de '
    'su propia agenda del día, y los administradores no pueden consultar '
    'rápidamente la ocupación de cada especialidad.',
    'Procesos redundantes: el personal administrativo dedica una cantidad '
    'significativa de tiempo a la búsqueda manual de horarios disponibles, '
    'confirmación telefónica y registro de cancelaciones.',
]

for p in problemas:
    doc.add_paragraph(p, style='List Bullet')

doc.add_paragraph(
    'Frente a esta problemática, surge la necesidad de contar con un sistema '
    'de escritorio que automatice la gestión integral de citas médicas, '
    'permitiendo que pacientes, médicos y administradores interactúen desde '
    'una misma plataforma con roles y permisos diferenciados, garantizando la '
    'integridad de los datos y la disponibilidad de la información en tiempo real.'
)

doc.add_page_break()

# ════════════════════════════════════════════
# 2. OBJETIVO GENERAL Y ESPECÍFICOS
# ════════════════════════════════════════════
add_title('2. Objetivo general y objetivos específicos')

add_title('2.1 Objetivo general', 12)
doc.add_paragraph(
    'Desarrollar una aplicación de escritorio para la gestión integral de citas '
    'médicas que permita automatizar los procesos de agendamiento, confirmación, '
    'atención y seguimiento clínico, integrando a pacientes, médicos y '
    'administradores en una plataforma unificada con roles diferenciados y '
    'validaciones automatizadas.'
)

add_title('2.2 Objetivos específicos', 12)

objetivos = [
    'Implementar un módulo de autenticación y registro de usuarios con '
    'validación de datos (email con verificación DNS, teléfono colombiano, '
    'formato de nombres y documentos) que garantice la integridad de la '
    'información ingresada.',
    'Diseñar un sistema de agendamiento de citas con detección dinámica de '
    'disponibilidad basado en la agenda del médico, bloqueos programados y '
    'duración variable del servicio.',
    'Desarrollar un mecanismo automático de detección de inasistencias que '
    'evalúe periódicamente las citas vencidas y las marque como "No asistió" '
    'sin intervención manual.',
    'Implementar un módulo de historial clínico que asocie diagnósticos, '
    'recetas y remisiones a cada cita completada, accesible tanto para el '
    'médico como para el paciente.',
    'Construir dashboards diferenciados por rol (paciente, médico, '
    'administrador) con operaciones CRUD completas sobre pacientes, médicos, '
    'citas y horarios.',
    'Incorporar validaciones de negocio como frecuencia máxima de citas por '
    'especialidad, sobrecupo controlado, reprogramación automática al eliminar '
    'médicos y bloqueo de agenda parcial o total.',
]

for o in objetivos:
    doc.add_paragraph(o, style='List Number')

doc.add_page_break()

# ════════════════════════════════════════════
# 3. JUSTIFICACIÓN
# ════════════════════════════════════════════
add_title('3. Justificación')

doc.add_paragraph(
    'La gestión manual de citas médicas representa un cuello de botella '
    'significativo en la operación diaria de las IPS. La implementación de '
    'Kairos busca resolver estos problemas mediante la automatización de los '
    'procesos clave, generando los siguientes beneficios:'
)

justificaciones = [
    ('Centralización de la información',
     'Todos los datos de pacientes, médicos, citas e historiales clínicos '
     'residen en una base de datos SQLite embebida, eliminando la dispersión '
     'de información en archivos físicos y hojas de cálculo.'),
    ('Reducción de errores operativos',
     'Las validaciones automáticas de formato (email, teléfono, documento, '
     'nombres) evitan el ingreso de datos inválidos. La detección de '
     'solapamiento de horarios mediante lógica de intervalos previene el '
     'doble agendamiento.'),
    ('Trazabilidad clínica',
     'Cada consulta completada deja un registro en el historial clínico con '
     'diagnóstico, receta y remisión, permitiendo al médico y al paciente '
     'consultar el historial completo de atenciones.'),
    ('Eficiencia operativa',
     'Los médicos visualizan su agenda del día filtrada por estado, los '
     'administradores gestionan horarios y bloqueos desde un solo panel, y '
     'los pacientes agendan sus citas sin intervención del personal '
     'administrativo.'),
    ('Reducción de inasistencias',
     'El sistema detecta automáticamente las citas no asistidas cada 5 '
     'minutos, y tanto el médico como el administrador pueden marcar '
     'inasistencias manualmente, generando datos precisos para la toma de '
     'decisiones.'),
    ('Arquitectura portable y de bajo costo',
     'Al usar SQLite como base de datos embebida, la aplicación no requiere '
     'instalación de servidores ni configuración de red. El archivo único '
     'de base de datos es fácil de respaldar y transportar.'),
    ('Escalabilidad académica',
     'La arquitectura en capas (modelo, DAO, servicio, controlador, vista) '
     'permite que el proyecto sea fácilmente extensible: agregar un nuevo '
     'rol, una nueva especialidad o migrar a MySQL requieren cambios '
     'localizados sin reescribir el sistema completo.'),
]

for tit, cuerpo in justificaciones:
    p = doc.add_paragraph()
    run = p.add_run(f'{tit}: ')
    run.bold = True
    p.add_run(cuerpo)

doc.add_paragraph(
    'En el contexto académico, el proyecto integra los cuatro pilares de la '
    'Programación Orientada a Objetos (herencia, polimorfismo, encapsulamiento '
    'y abstracción), estructuras de datos (Stack, Queue), manejo de excepciones, '
    'bases de datos relacionales con JDBC y SQLite, interfaces gráficas con '
    'JavaFX y FXML, así como patrones de diseño como Singleton, DAO, MVC, '
    'Template Method y State, consolidando los conocimientos adquiridos en el '
    'curso.'
)

doc.add_page_break()

# ════════════════════════════════════════════
# 4. DIAGRAMA DE CLASES UML
# ════════════════════════════════════════════
add_title('4. Diagrama de clases UML')

doc.add_paragraph(
    'A continuación se presenta el diagrama de clases UML del sistema. '
    'El código está escrito en sintaxis Mermaid; puede copiarse y pegarse en '
    'un editor Mermaid (como mermaid.live o la extensión Mermaid de VS Code) '
    'para visualizarlo o exportarlo como imagen.'
)

# Mermaid code as code block
add_title('Código Mermaid del diagrama', 12)

mermaid_code = """classDiagram
    class Usuario {
        <<abstract>>
        -id String
        -nombre String
        -email String
        -password String
        -rol Rol
        +getMenuOpciones() String[]*
        +login(String, String) boolean
    }

    class Paciente {
        -tipoDocumento String
        -numeroDocumento String
        -eps String
        -historialCitas Stack~Cita~
        +getMenuOpciones() String[]
    }

    class Medico {
        -especialidad Especialidad
        -agendaDelDia Queue~Cita~
        +getMenuOpciones() String[]
    }

    class Administrador {
        -codigoAdmin String
        +getMenuOpciones() String[]
    }

    Usuario <|-- Paciente
    Usuario <|-- Medico
    Usuario <|-- Administrador

    class Cita {
        -fecha LocalDate
        -horaInicio LocalTime
        -duracion int
        -estado EstadoCita
        -sobrecupo boolean
        +confirmar() boolean
        +cancelar() boolean
        +completar() boolean
    }

    class EstadoCita {
        <<enumeration>>
        PENDIENTE
        CONFIRMADA
        COMPLETADA
        CANCELADA
        NO_ASISTIO
    }

    class HistorialClinico {
        -diagnostico String
        -receta String
        -remision String
    }

    class AgendaMedica {
        -diaSemana int
        -horaInicio LocalTime
        -horaFin LocalTime
        -slotMinutos int
    }

    Cita --> EstadoCita
    Cita --> Paciente
    Cita --> Medico
    Cita "1" --> "0..1" HistorialClinico
    Medico "1" --> "0..*" AgendaMedica

    package "Vista (FXML)" {
        class login_fxml
        class dashboardPaciente_fxml
        class dashboardMedico_fxml
        class dashboardAdmin_fxml
    }

    package "Controlador" {
        class LoginController
        class DashboardPacienteController
        class DashboardMedicoController
        class DashboardAdminController
    }

    package "Servicio (Reglas de negocio)" {
        class CitaService
        class DisponibilidadService
        class Session
    }

    package "DAO (Persistencia)" {
        class CitaDAO
        class UsuarioDAO
    }

    LoginController --> Session
    LoginController --> UsuarioDAO
    DashboardPacienteController --> CitaService
    DashboardMedicoController --> CitaService
    DashboardAdminController --> CitaService
    CitaService --> CitaDAO
    CitaService --> DisponibilidadService"""

# Write Mermaid to a code block-style
from docx.shared import RGBColor

p = doc.add_paragraph()
run = p.add_run(mermaid_code)
run.font.name = 'Consolas'
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

doc.add_paragraph(
    'Instrucciones: copie el código anterior y péguelo en un editor Mermaid '
    '(https://mermaid.live o extensión Mermaid en VS Code) para generar la '
    'imagen del diagrama.'
)

doc.add_page_break()

# ════════════════════════════════════════════
# 5. MANUAL DE USUARIO
# ════════════════════════════════════════════
add_title('5. Manual de usuario')

# 5.1 Requisitos
add_title('5.1 Requisitos del sistema', 12)

doc.add_paragraph('Para ejecutar Kairos se necesita:')

reqs_text = [
    'Java 17 o superior (JDK 17+),',
    'Maven 3.8 o superior (para compilar desde consola),',
    'NetBeans 17+ (recomendado) o cualquier IDE con soporte Maven,',
    'Conexión a Internet solo para la validación DNS de email (opcional),',
    'Sistema operativo: Windows, macOS o Linux.',
]
for r in reqs_text:
    doc.add_paragraph(r, style='List Bullet')

# 5.2 Instalación
add_title('5.2 Instalación y ejecución', 12)

doc.add_paragraph('Opción 1 — Desde Maven:')
doc.add_paragraph('    mvn clean javafx:run', style='List Bullet')
doc.add_paragraph('Opción 2 — Desde NetBeans:')
doc.add_paragraph('    Abrir el proyecto y presionar F6 (Run).', style='List Bullet')
doc.add_paragraph('Opción 3 — JAR ejecutable:')
doc.add_paragraph('    mvn clean package assembly:single', style='List Bullet')
doc.add_paragraph('    java -jar target/SistemaCitasMedicas-1.0-SNAPSHOT-jar-with-dependencies.jar', style='List Bullet')

doc.add_paragraph(
    'En la primera ejecución, el sistema crea automáticamente la base de datos '
    'SQLite (citasmedicas.db) con datos de demostración precargados.'
)

# 5.3 Login
add_title('5.3 Inicio de sesión', 12)

doc.add_paragraph(
    'Al ejecutar la aplicación aparece la pantalla de inicio de sesión. '
    'El usuario debe ingresar su correo electrónico y contraseña registrados '
    'y hacer clic en "Iniciar sesión". El sistema redirige automáticamente al '
    'dashboard correspondiente según el rol del usuario.'
)

doc.add_paragraph('Credenciales de demostración:')

# Tabla de credenciales
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Rol', 'Email', 'Contraseña']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

data = [
    ['Paciente', 'paciente@demo.com', '1234'],
    ['Médico', 'medico@demo.com', '1234'],
    ['Administrador', 'admin@demo.com', '1234'],
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i + 1].cells[j].text = val

doc.add_paragraph()
doc.add_paragraph(
    'Si las credenciales son incorrectas, el sistema muestra un mensaje de '
    'error en la pantalla. El usuario puede hacer clic en "Registrarse" si '
    'aún no tiene una cuenta de paciente.'
)

doc.add_paragraph('(Insertar aquí captura de pantalla de la ventana de login)')

# 5.4 Registro
add_title('5.4 Registro de paciente', 12)

doc.add_paragraph(
    'El registro se realiza mediante un asistente de dos pasos:'
)

doc.add_paragraph('Paso 1: Datos básicos', style='List Bullet')
doc.add_paragraph(
    '    El usuario ingresa nombre, apellido y selecciona el rol '
    '("Paciente", "Médico" o "Administrador"). Si selecciona '
    '"Administrador", debe ingresar el código de administrador '
    'validado contra la base de datos.',
    style='List Bullet'
)
doc.add_paragraph('Paso 2: Datos específicos del rol', style='List Bullet')
doc.add_paragraph(
    '    Para un paciente: email, contraseña (con doble confirmación), '
    'teléfono móvil colombiano, tipo de documento (ComboBox con CC, TI, '
    'CE, Pasaporte), número de documento, fecha de nacimiento, dirección '
    'y EPS.',
    style='List Bullet'
)

doc.add_paragraph('Validaciones aplicadas durante el registro:')

validaciones = [
    'Email: se verifica el formato (usuario@dominio.ext) Y se consulta '
    'DNS para confirmar que el dominio tiene registros MX (puede recibir '
    'correos).',
    'Teléfono: debe ser un número móvil colombiano válido de 10 dígitos '
    'que empiece con 3 (opcionalmente con prefijo +57).',
    'Nombre y apellido: solo letras (incluyendo tildes y ñ) y espacios. '
    'Números y caracteres especiales son rechazados.',
    'Número de documento: solo dígitos.',
    'EPS: solo letras y espacios.',
    'Confirmación de contraseña: ambos campos deben coincidir.',
    'Email duplicado: no se permite registrar dos usuarios con el mismo email.',
]
for v in validaciones:
    doc.add_paragraph(v, style='List Bullet')

doc.add_paragraph('(Insertar aquí capturas de pantalla del registro: paso 1 y paso 2)')

# 5.5 Dashboard Paciente
add_title('5.5 Dashboard del paciente', 12)

doc.add_paragraph(
    'Después de iniciar sesión como paciente, se muestra el dashboard con '
    'las siguientes opciones en el menú lateral:'
)

opciones_paciente = [
    'Solicitar cita: permite agendar una nueva cita médica. El paciente '
    'selecciona especialidad, servicio, médico, fecha, hora (solo se muestran '
    'las disponibles), tipo (presencial/virtual) y motivo de consulta '
    '(TextArea expandible con vista completa en diálogo modal).',
    'Mis citas: muestra las citas activas del paciente en una tabla con '
    'filas coloreadas por estado (PENDIENTE, CONFIRMADA, COMPLETADA, etc.). '
    'El paciente puede cancelar una cita si aún no ha sido completada.',
    'Modificar datos: permite actualizar la información personal del paciente '
    '(email, teléfono, dirección, EPS, etc.) con las mismas validaciones del '
    'registro.',
    'Cerrar sesión: vuelve a la pantalla de login.',
]
for o in opciones_paciente:
    doc.add_paragraph(o, style='List Bullet')

doc.add_paragraph(
    'La tabla de citas incluye barra de búsqueda, tooltips sobre cada '
    'fila y un contador que muestra el total de citas del paciente. '
    'Las citas canceladas se excluyen de la vista del paciente.'
)

doc.add_paragraph('(Insertar aquí captura de pantalla del dashboard del paciente)')

# 5.6 Dashboard Médico
add_title('5.6 Dashboard del médico', 12)

doc.add_paragraph(
    'El médico cuenta con las siguientes funcionalidades:'
)

opciones_medico = [
    'Agenda del día: lista de citas del día actual filtradas por estado '
    '(PENDIENTE y CONFIRMADA), ordenadas por hora. Incluye barra de '
    'búsqueda y contador de citas.',
    'Ver siguiente paciente: atiende al primer paciente en la cola (FIFO), '
    'abriendo automáticamente la ventana de consulta.',
    'Iniciar consulta: abre el formulario de historial clínico donde el '
    'médico ingresa diagnóstico (obligatorio), enfermedad actual, receta, '
    'remisión y notas. Al guardar, la cita pasa a estado COMPLETADA.',
    'Marcar como no asistió: cambia el estado de la cita a NO_ASISTIO '
    'sin necesidad de llenar historial clínico.',
    'Agendar control: agenda una cita de seguimiento (origen=CONTROL) '
    'para un paciente, con duración reducida según el servicio.',
    'Reprogramar cita: permite cambiar la fecha y hora de una cita '
    'existente a un nuevo slot disponible.',
    'Gestionar horarios: administra su propia agenda médica (días de la '
    'semana, horarios de atención, duración de slots).',
    'Calendario visual: integración con CalendarFX que muestra las citas '
    'del mes en formato de calendario.',
    'Consultar historial del paciente: antes de atender, el médico puede '
    'revisar el historial clínico completo del paciente.',
]
for o in opciones_medico:
    doc.add_paragraph(o, style='List Bullet')

doc.add_paragraph(
    'Si una cita está en estado PENDIENTE y el médico la atiende, '
    'el sistema la confirma automáticamente y luego la completa en '
    'un solo paso.'
)

doc.add_paragraph('(Insertar aquí captura de pantalla del dashboard del médico)')

# 5.7 Dashboard Admin
add_title('5.7 Dashboard del administrador', 12)

doc.add_paragraph(
    'El administrador tiene acceso completo a la gestión del sistema:'
)

opciones_admin = [
    'Gestionar pacientes: lista todos los pacientes activos con opciones '
    'de editar datos y desactivar (borrado lógico).',
    'Gestionar médicos: lista los médicos activos, permite registrar '
    'nuevos médicos con todos sus datos (incluyendo registro médico y '
    'especialidad) y desactivar médicos existentes.',
    'Gestionar citas: visualiza todas las citas del sistema con filtros '
    'por estado, médico y fecha. Permite agregar nuevas citas (incluyendo '
    'sobrecupo), modificar citas existentes y cancelar citas.',
    'Gestionar horarios: administra la agenda médica de cualquier médico '
    'del sistema (días, horarios y slots).',
    'Gestionar bloqueos: permite bloquear parcial o totalmente la agenda '
    'de un médico en una fecha específica (ej: vacaciones, junta médica).',
    'Marcar como no asistió: similar al médico, para cualquier cita '
    'del sistema.',
]
for o in opciones_admin:
    doc.add_paragraph(o, style='List Bullet')

doc.add_paragraph(
    'Todas las tablas del administrador incluyen barras de búsqueda, '
    'tooltips y contadores. La interfaz usa pestañas (TabPane) para '
    'organizar las secciones.'
)

doc.add_paragraph('(Insertar aquí captura de pantalla del dashboard del administrador)')

# 5.8 Funcionalidades comunes
add_title('5.8 Funcionalidades comunes', 12)

doc.add_paragraph(
    'Independientemente del rol, el sistema ofrece las siguientes '
    'funcionalidades transversales:'
)

comunes = [
    'Filas coloreadas por estado: las citas en la tabla se muestran con '
    'colores según su estado (verde para CONFIRMADA, rojo para CANCELADA, '
    'gris para COMPLETADA, etc.), facilitando la identificación visual.',
    'Búsqueda en tablas: todas las tablas del sistema incluyen una barra '
    'de búsqueda que filtra en tiempo real los resultados.',
    'Tooltips: al pasar el mouse sobre una fila, se muestra un tooltip '
    'con información detallada de la cita.',
    'Contadores: cada sección muestra el número total de elementos '
    '(ej: "9 citas encontradas").',
    'Diálogos de confirmación: las acciones destructivas (cancelar cita, '
    'desactivar usuario) muestran un cuadro de diálogo de confirmación '
    'antes de ejecutarse.',
    'Notificaciones de error: los mensajes de error se muestran en la '
    'interfaz (labels de color rojo) sin interrumpir el flujo del usuario.',
    'Detección automática de inasistencias: cada 5 minutos el sistema '
    'verifica y marca automáticamente las citas vencidas como NO_ASISTIO.',
]
for c in comunes:
    doc.add_paragraph(c, style='List Bullet')

# Guardar
outdir = os.path.dirname(__file__)
tmp_path = os.path.join(outdir, '_temp_kairos.docx')
output_path = os.path.join(outdir, 'Documento_Proyecto_Kairos.docx')
doc.save(tmp_path)
if os.path.exists(output_path):
    os.remove(output_path)
os.rename(tmp_path, output_path)
doc.save(output_path)
print(f'Documento generado: {output_path}')
